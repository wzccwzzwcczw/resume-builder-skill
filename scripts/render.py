#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""通用简历渲染器：一份 resume.json → HTML / PDF / Word 三格式。
用法：
  python3 render.py resume.json --out ~/Desktop --formats html,pdf,docx --style visual
  --style visual : 蓝白视觉版（投牛客站内/发人看/打印）
  --style ats    : ATS 安全版（单栏纯净，投大厂官网/邮箱）
默认段落顺序（校招）：基本信息 · 教育 · 实习 · 项目 · 技能 · 荣誉
"""
import json, sys, os, re, argparse, subprocess, shutil, html as _html

def find_chrome():
    """跨平台定位 Chrome/Chromium：环境变量 > PATH > 常见安装路径。"""
    env = os.environ.get("RESUME_CHROME") or os.environ.get("CHROME_PATH")
    if env and os.path.exists(env):
        return env
    for name in ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser", "chrome"):
        p = shutil.which(name)
        if p:
            return p
    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",   # macOS
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/usr/bin/google-chrome", "/usr/bin/chromium", "/usr/bin/chromium-browser",  # Linux
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",          # Windows
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    ]
    for c in candidates:
        if os.path.exists(c):
            return c
    return None
DEFAULT_ORDER = ["basics", "education", "internships", "projects", "skills", "honors"]
SECTION_TITLE = {"education": "教育背景", "internships": "实习经历",
                 "projects": "项目经历", "skills": "专业技能", "honors": "荣誉奖项"}
SECTION_EN = {"education": "EDUCATION", "internships": "INTERNSHIP",
              "projects": "PROJECTS", "skills": "SKILLS", "honors": "HONORS"}

def esc(s): return _html.escape(str(s)) if s not in (None, "") else ""
def md_b(s):  # **加粗** → <b>
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", esc(s))

# ---------------- HTML ----------------
def build_html(d, style="visual"):
    b = d.get("basics", {})
    order = d.get("meta", {}).get("section_order") or DEFAULT_ORDER
    ic = {"phone": "📱", "email": "✉️", "github": "🐙 GitHub：", "portfolio": "🔗 ", "edu": "🎓 "} if style == "visual" else \
         {"phone": "电话：", "email": "邮箱：", "github": "GitHub：", "portfolio": "", "edu": ""}
    contact = [x for x in [
        (ic["phone"], b.get("phone")), (ic["email"], b.get("email")),
        (ic["github"], b.get("github")), (ic["portfolio"], b.get("portfolio"))] if x[1]]
    edu_line = ""
    if d.get("education"):
        e0 = d["education"][0]
        edu_line = " · ".join([x for x in [e0.get("school"), e0.get("major"),
                    e0.get("degree"), e0.get("grad_year")] if x])

    def sec_education():
        rows = ""
        for e in d.get("education", []):
            head = " · ".join([x for x in [e.get("school"), e.get("major"), e.get("degree")] if x])
            rows += f'<div class="row"><div class="t">{esc(head)}</div><div class="d">{esc(e.get("date",""))}</div></div>'
            sub = " ｜ ".join([x for x in [f'GPA {e["gpa"]}' if e.get("gpa") else "",
                   f'相关课程：{e["courses"]}' if e.get("courses") else ""] if x])
            if sub: rows += f'<div class="sub">{esc(sub)}</div>'
        return rows

    def sec_internships():
        out = ""
        for it in d.get("internships", []):
            out += f'<div class="row"><div class="t">{esc(it.get("company",""))} · {esc(it.get("title",""))}</div><div class="d">{esc(it.get("date",""))}</div></div>'
            if it.get("bullets"):
                out += "<ul>" + "".join(f"<li>{md_b(x)}</li>" for x in it["bullets"]) + "</ul>"
        return out

    def sec_projects():
        out = ""
        for p in d.get("projects", []):
            tag = f' <small>{esc(p.get("tag"))}</small>' if p.get("tag") else ""
            out += f'<div class="row"><div class="t">{esc(p.get("name",""))}{tag}</div><div class="d">{esc(p.get("date",""))}</div></div>'
            meta = " ｜ ".join([x for x in [f'技术栈：{p["stack"]}' if p.get("stack") else "",
                    f'角色：{p["role"]}' if p.get("role") else ""] if x])
            if meta: out += f'<div class="sub">{esc(meta)}</div>'
            if p.get("bullets"):
                out += "<ul>" + "".join(f"<li>{md_b(x)}</li>" for x in p["bullets"]) + "</ul>"
        return out

    def sec_skills():
        rows = ""
        for s in d.get("skills", []):
            rows += f'<span class="k">{esc(s.get("category",""))}</span><span class="v">{esc(s.get("items",""))}</span>'
        return f'<div class="skills">{rows}</div>'

    def sec_honors():
        out = "<ul>"
        for h in d.get("honors", []):
            tag = f'<span class="tag">{esc(h.get("type"))}</span>' if h.get("type") else ""
            out += f'<li>{tag}{md_b(h.get("text",""))}</li>'
        return out + "</ul>"

    builders = {"education": sec_education, "internships": sec_internships,
                "projects": sec_projects, "skills": sec_skills, "honors": sec_honors}
    body = ""
    for key in order:
        if key == "basics":
            continue
        inner = builders[key]()
        if not inner.strip() or inner in ("<ul></ul>", '<div class="skills"></div>'):
            continue
        body += (f'<section><h2><span class="bar"></span>{SECTION_TITLE[key]}'
                 f'<span class="en">{SECTION_EN[key]}</span></h2>{inner}</section>')

    contact_html = "".join(f'<span>{esc(ic)}<b>{esc(v)}</b></span>' for ic, v in contact)
    role = f'<div class="role">求职意向：{esc(b.get("role"))}</div>' if b.get("role") else ""
    eduline = f'<span>{ic["edu"]}<b>{esc(edu_line)}</b></span>' if edu_line else ""

    css_visual = """
  :root{--blue:#2563eb;--blue-d:#1e40af;--ink:#1f2329;--gray:#6b7280;--line:#e5e7eb;--bg:#eff4ff;}
  body{font-family:-apple-system,"PingFang SC","Microsoft YaHei",sans-serif;color:var(--ink);font-size:10.2px;line-height:1.45;}
  .header{background:linear-gradient(120deg,#2563eb,#1e3a8a);color:#fff;padding:16px 30px 13px;position:relative;}
  .header::after{content:"";position:absolute;left:0;bottom:0;width:100%;height:4px;background:linear-gradient(90deg,#60a5fa,#93c5fd);}
  .name{font-size:24px;font-weight:800;letter-spacing:1px;}
  .role{display:inline-block;margin-top:4px;font-size:12px;font-weight:600;color:#dbeafe;border:1px solid rgba(255,255,255,.5);padding:2px 12px;border-radius:20px;}
  .contact{margin-top:9px;font-size:10px;color:#e7eefc;display:flex;flex-wrap:wrap;gap:4px 18px;}
  .contact b{color:#fff;font-weight:600;}
  .body{padding:10px 30px 14px;}
  section{margin-top:9px;}
  h2{font-size:12.5px;color:var(--blue);font-weight:800;letter-spacing:.5px;padding-bottom:3px;border-bottom:2px solid var(--blue);display:flex;align-items:center;gap:7px;}
  h2 .bar{display:inline-block;width:4px;height:12px;background:var(--blue);border-radius:2px;}
  h2 .en{color:#9db8e8;font-size:9px;font-weight:600;margin-left:auto;letter-spacing:1px;}
  .row{display:flex;justify-content:space-between;align-items:baseline;margin-top:6px;}
  .row .t{font-size:11.3px;font-weight:700;} .row .t small{font-weight:600;color:var(--blue);}
  .row .d{font-size:9.5px;color:var(--gray);white-space:nowrap;}
  .sub{font-size:9.7px;color:var(--gray);margin-top:1px;}
  ul{list-style:none;margin:3px 0 0;} li{position:relative;padding-left:13px;margin:2px 0;font-size:10px;}
  li::before{content:"";position:absolute;left:2px;top:6px;width:5px;height:5px;border-radius:50%;background:var(--blue);}
  li b{color:var(--blue-d);font-weight:700;}
  .skills{margin-top:7px;display:grid;grid-template-columns:max-content 1fr;gap:5px 12px;font-size:10.1px;}
  .skills .k{font-weight:700;color:var(--blue-d);background:var(--bg);padding:2px 9px;border-radius:5px;text-align:center;white-space:nowrap;}
  .skills .v{align-self:center;}
  .tag{display:inline-block;background:var(--bg);color:var(--blue-d);font-size:9.3px;font-weight:600;padding:1px 8px;border-radius:5px;margin-right:5px;}
"""
    css_ats = """
  body{font-family:"Times New Roman","SimSun",serif;color:#000;font-size:11px;line-height:1.5;}
  .header{padding:18px 26px 6px;text-align:center;}
  .name{font-size:22px;font-weight:bold;} .role{font-size:12px;margin-top:3px;}
  .contact{margin-top:6px;font-size:10.5px;display:block;} .contact span{margin:0 8px;}
  .contact b{font-weight:normal;}
  .body{padding:6px 26px 18px;}
  section{margin-top:11px;}
  h2{font-size:13px;font-weight:bold;border-bottom:1px solid #000;padding-bottom:2px;text-transform:none;}
  h2 .bar,h2 .en{display:none;}
  .row{display:flex;justify-content:space-between;margin-top:7px;} .row .t{font-weight:bold;font-size:11px;}
  .row .t small{font-weight:normal;} .row .d{font-size:10px;white-space:nowrap;}
  .sub{font-size:10px;margin-top:1px;}
  ul{margin:3px 0 0 18px;} li{margin:2px 0;font-size:10.5px;} li b{font-weight:bold;}
  .skills{margin-top:5px;} .skills .k{font-weight:bold;} .skills .k::after{content:"：";}
  .skills{display:grid;grid-template-columns:max-content 1fr;gap:3px 8px;font-size:10.5px;}
  .tag{font-weight:bold;} .tag::after{content:" ";}
"""
    css = css_visual if style == "visual" else css_ats
    return f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">
<title>{esc(b.get('name',''))}的简历</title>
<style>@page{{size:A4;margin:0;}}*{{margin:0;padding:0;box-sizing:border-box;-webkit-print-color-adjust:exact;print-color-adjust:exact;}}
.page{{width:210mm;min-height:297mm;background:#fff;}}{css}</style></head>
<body><div class="page">
<div class="header"><div class="name">{esc(b.get('name',''))}</div>{role}
<div class="contact">{contact_html}{eduline}</div></div>
<div class="body">{body}</div></div></body></html>"""

# ---------------- PDF ----------------
def render_pdf(html_path, pdf_path):
    chrome = find_chrome()
    if not chrome:
        print("⚠️  未找到 Chrome/Chromium，跳过 PDF。请安装 Chrome，或设置环境变量 "
              "RESUME_CHROME=/path/to/chrome（HTML/Word 不受影响）", file=sys.stderr)
        return False
    subprocess.run([chrome, "--headless", "--disable-gpu", "--no-sandbox",
        "--no-pdf-header-footer", f"--print-to-pdf={pdf_path}",
        "--virtual-time-budget=3000", f"file://{html_path}"],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return os.path.exists(pdf_path)

# ---------------- Word ----------------
def build_docx(d, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    BLUE = RGBColor(0x25, 0x63, 0xeb)
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'PingFang SC'; st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
    def f(r): r.font.name = 'PingFang SC'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
    def runs(p, text):  # support **bold**
        for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", str(text))):
            r = p.add_run(seg); r.bold = (i % 2 == 1); f(r)
    b = d.get("basics", {})
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(b.get("name", "")); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE; f(r)
    if b.get("role"):
        rp = doc.add_paragraph(); rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = rp.add_run("求职意向：" + b["role"]); rr.font.size = Pt(11); f(rr)
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = " | ".join([x for x in [b.get("phone"), b.get("email"), b.get("github"), b.get("portfolio")] if x])
    rr = cp.add_run(contact); rr.font.size = Pt(9.5); f(rr)

    def heading(t):
        p = doc.add_heading(level=1); r = p.add_run(t); r.font.color.rgb = BLUE; f(r)
    order = d.get("meta", {}).get("section_order") or DEFAULT_ORDER
    for key in order:
        if key == "basics": continue
        if key == "education" and d.get("education"):
            heading("教育背景")
            for e in d["education"]:
                p = doc.add_paragraph(); rr = p.add_run(" · ".join([x for x in [e.get("school"), e.get("major"), e.get("degree")] if x]) + ("    " + e.get("date","") if e.get("date") else "")); rr.bold = True; f(rr)
                sub = " ｜ ".join([x for x in [f'GPA {e["gpa"]}' if e.get("gpa") else "", f'相关课程：{e["courses"]}' if e.get("courses") else ""] if x])
                if sub: rr = doc.add_paragraph().add_run(sub); rr.font.size = Pt(9.5); f(rr)
        elif key == "internships" and d.get("internships"):
            heading("实习经历")
            for it in d["internships"]:
                p = doc.add_paragraph(); rr = p.add_run(f'{it.get("company","")} · {it.get("title","")}    {it.get("date","")}'); rr.bold = True; f(rr)
                for x in it.get("bullets", []):
                    bp = doc.add_paragraph(style='List Bullet'); runs(bp, x)
        elif key == "projects" and d.get("projects"):
            heading("项目经历")
            for pj in d["projects"]:
                p = doc.add_paragraph(); rr = p.add_run(f'{pj.get("name","")} {pj.get("tag","")}    {pj.get("date","")}'); rr.bold = True; f(rr)
                meta = " ｜ ".join([x for x in [f'技术栈：{pj["stack"]}' if pj.get("stack") else "", f'角色：{pj["role"]}' if pj.get("role") else ""] if x])
                if meta: rr = doc.add_paragraph().add_run(meta); rr.font.size = Pt(9.5); f(rr)
                for x in pj.get("bullets", []):
                    bp = doc.add_paragraph(style='List Bullet'); runs(bp, x)
        elif key == "skills" and d.get("skills"):
            heading("专业技能")
            for s in d["skills"]:
                p = doc.add_paragraph(); rr = p.add_run(s.get("category","") + "："); rr.bold = True; f(rr); rr2 = p.add_run(s.get("items","")); f(rr2)
        elif key == "honors" and d.get("honors"):
            heading("荣誉奖项")
            for hn in d["honors"]:
                bp = doc.add_paragraph(style='List Bullet'); runs(bp, (f'【{hn["type"]}】' if hn.get("type") else "") + hn.get("text",""))
    doc.save(path)

# ---------------- main ----------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("json"); ap.add_argument("--out", default=os.path.expanduser("~/Desktop"))
    ap.add_argument("--formats", default="html,pdf,docx")
    ap.add_argument("--style", default="visual", choices=["visual", "ats"])
    a = ap.parse_args()
    d = json.load(open(a.json, encoding="utf-8"))
    name = d.get("basics", {}).get("name", "简历")
    fmts = [x.strip() for x in a.formats.split(",")]
    os.makedirs(a.out, exist_ok=True)
    base = os.path.join(a.out, f"简历_{name}")
    made = []
    if "html" in fmts or "pdf" in fmts:
        html_doc = build_html(d, a.style)
        html_path = base + (f"_{a.style}.html")
        open(html_path, "w", encoding="utf-8").write(html_doc)
        if "html" in fmts: made.append(html_path)
        if "pdf" in fmts:
            pdf_path = base + f"_{a.style}.pdf"   # 文件名带 style，visual/ats 不互相覆盖
            if render_pdf(html_path, pdf_path): made.append(pdf_path)
            if "html" not in fmts: os.remove(html_path)
    if "docx" in fmts:
        docx_path = base + f"_{a.style}.docx"
        build_docx(d, docx_path); made.append(docx_path)
    for m in made: print("✅", m)

if __name__ == "__main__":
    main()
